import sys
import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_custom_table(table, col_widths, col_alignments, header_bg="1E3A8A", alt_bg="F8FAFC"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        bg_color = header_bg if is_header else (alt_bg if row_idx % 2 == 1 else "FFFFFF")
        
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if is_header:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[col_idx])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            for p in cell.paragraphs:
                p.alignment = col_alignments[col_idx]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = 'Cordia New'
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(11)
                    else:
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = RGBColor(30, 41, 59)

def add_callout(doc, text_list, title="บทวิเคราะห์เชิงลึก (Data Engineering Insights)", bg_hex="EFF6FF", border_hex="3B82F6"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(30, 58, 138)
    
    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        r = p_item.add_run(f"• {item}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def generate_summary_and_answers_document():
    doc = Document()
    
    # Page setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Cordia New'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(30, 41, 59)
    
    # Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("รายงานสรุปผลการทดลองและคำตอบ (Lab Report & Solutions)")
    run_title.font.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(30, 58, 138)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Lab Week 8: Data Integration Pipeline — TechTrove E-Commerce\nรายวิชา: Data Warehousing Concepts and Design")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    # -------------------------------------------------------------------------
    # 1. รายงานสรุปผลภาพรวม (Executive Summary & Key Metrics)
    # -------------------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. รายงานสรุปผลภาพรวมของไปป์ไลน์ (Pipeline Execution Summary)")
    r_h1.font.bold = True
    r_h1.font.size = Pt(13.5)
    r_h1.font.color.rgb = RGBColor(30, 58, 138)
    
    p_intro = doc.add_paragraph(
        "กระบวนการ Data Integration Pipeline ได้ทำการรวบรวมข้อมูลจาก 4 ระบบ (คำสั่งซื้อรายเดือน ม.ค. และ ก.พ., "
        "ลูกค้า CRM, สินค้า Product Master และการชำระเงิน Payment Gateway) ผ่านการทำ Schema Alignment, "
        "Data Cleansing, การกำจัดข้อมูลซ้ำ (Deduplication), การตรวจสอบ Referential Integrity และการบังคับใช้ Business Rules "
        "จนได้ชุดข้อมูล Fact และ Dimension สำหรับ Data Warehouse ที่มีความถูกต้องสมบูรณ์ 100%"
    )
    p_intro.paragraph_format.space_after = Pt(6)
    
    # Metrics Table
    tbl_sum = doc.add_table(rows=6, cols=2)
    sum_data = [
        ("ตัวชี้วัดหลักของการประมวลผล (Key Metric)", "ผลลัพธ์ที่ได้จากการรันจริง (Verified Value)"),
        ("จำนวนคำสั่งซื้อดิบทั้งหมด (Raw Combined Orders)", "752 แถว (ม.ค. 361 แถว + ก.พ. 391 แถว)"),
        ("จำนวนคำสั่งซื้อหลังตัดข้อมูลซ้ำ (Deduplicated Orders)", "750 แถว (ตัดแถวซ้ำออก 2 แถว)"),
        ("จำนวนรายการขายที่ถูกต้องและชำระเงินสำเร็จ (fact_sales)", "660 ธุรกรรม (อัตราผ่านเกณฑ์ 87.77%)"),
        ("ยอดขายสุทธิรวมทั้งสิ้น (Total Net Sales Revenue)", "฿10,224,044.09 บาท"),
        ("จำนวนเหตุการณ์ Data Quality ที่บันทึกตรวจสอบ (DQ Log)", "96 รายการ (บันทึกใน data_quality_report.csv)"),
    ]
    for r_idx, row_vals in enumerate(sum_data):
        for c_idx, val in enumerate(row_vals):
            tbl_sum.cell(r_idx, c_idx).paragraphs[0].text = val
    format_custom_table(tbl_sum, [3.2, 3.3], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # -------------------------------------------------------------------------
    # 2. รายงานคุณภาพข้อมูลและ Data Quality Funnel
    # -------------------------------------------------------------------------
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. รายงานคุณภาพข้อมูลและขั้นตอนการกลั่นกรอง (Data Quality Funnel)")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13.5)
    r_h2.font.color.rgb = RGBColor(30, 58, 138)
    
    tbl_funnel = doc.add_table(rows=6, cols=4)
    funnel_data = [
        ("ขั้นตอน (Funnel Stage)", "จำนวนคงเหลือ (Rows)", "จำนวนที่คัดออก (Filtered)", "เหตุผลและความผิดปกติที่ตรวจพบ"),
        ("1. Raw Orders Extracted", "752", "-", "คำสั่งซื้อดิบ ม.ค. (361) + ก.พ. (391)"),
        ("2. Deduplicated Orders", "750", "2", "ลบคำสั่งซื้อซ้ำ (ORD000056, ORD000416)"),
        ("3. Referential Integrity Check", "726", "24", "ไม่พบใน CRM 22 แถว, ไม่พบใน Product Master 2 แถว"),
        ("4. Business Rules Validation", "724", "2", "unit_price เป็น Null 2 แถว, quantity <= 0 อีก 2 แถว"),
        ("5. Successful Paid Sales (fact_sales)", "660", "64", "สถานะชำระเงินไม่สำเร็จ (FAILED 46 รายการ, REFUNDED 18 รายการ)")
    ]
    for r_idx, row_vals in enumerate(funnel_data):
        for c_idx, val in enumerate(row_vals):
            tbl_funnel.cell(r_idx, c_idx).paragraphs[0].text = val
    format_custom_table(tbl_funnel, [2.2, 1.1, 1.2, 2.0], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # -------------------------------------------------------------------------
    # 3. คำตอบคำถามวิเคราะห์เชิงธุรกิจ (Answers to Analysis Questions 1-6)
    # -------------------------------------------------------------------------
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. คำตอบคำถามวิเคราะห์เชิงธุรกิจ (Answers to Business Analysis Questions)")
    r_h3.font.bold = True
    r_h3.font.size = Pt(13.5)
    r_h3.font.color.rgb = RGBColor(30, 58, 138)
    
    # Q1
    p1 = doc.add_paragraph()
    p1.add_run("คำถามข้อที่ 1: หลังรวมไฟล์ orders มีจำนวนแถวเท่าใด และเหลือกี่แถวหลังลบ duplicate?\n").font.bold = True
    p1.add_run(
        "• คำตอบ:\n"
        "  - หลังรวมไฟล์คำสั่งซื้อ ม.ค. (361 แถว) และ ก.พ. (391 แถว) มีจำนวนแถวทั้งสิ้น 752 แถว\n"
        "  - เมื่อตัดแถวที่ซ้ำซ้อนออกด้วยกฎ subset=['order_id'], keep='last' มีแถวซ้ำถูกลบ 2 แถว (ได้แก่ ORD000056 และ ORD000416)\n"
        "  - คงเหลือคำสั่งซื้อที่ไม่ซ้ำซ้อนทั้งสิ้น 750 แถว"
    )
    
    # Q2
    p2 = doc.add_paragraph()
    p2.add_run("คำถามข้อที่ 2: มีแถวที่ customer_id หรือ product_id ไม่พบใน Master Data อย่างละกี่แถว?\n").font.bold = True
    p2.add_run(
        "• คำตอบ:\n"
        "  - customer_id ที่ไม่พบใน CRM Master Data: มีจำนวน 22 แถว (ได้แก่รหัส C0161, C0162, C0163, C0164, C0165 ซึ่งเป็นลูกค้านอกระบบ)\n"
        "  - product_id ที่ไม่พบใน Product Master: มีจำนวน 2 แถว (ได้แก่รหัสสินค้า P999 ในคำสั่งซื้อ ORD000022 และ ORD000382)"
    )
    
    # Q3
    p3 = doc.add_paragraph()
    p3.add_run("คำถามข้อที่ 3: มียอดขายที่ใช้ได้จริงกี่ธุรกรรม และยอดขายสุทธิรวมเท่าใด?\n").font.bold = True
    p3.add_run(
        "• คำตอบ:\n"
        "  - จำนวนธุรกรรมยอดขายที่ใช้ได้จริงใน Fact Sales: 660 ธุรกรรม (คิดเป็น 87.77% ของคำสั่งซื้อ)\n"
        "  - ยอดขายสุทธิรวมทั้งสิ้น (Total Net Sales Revenue): ฿10,224,044.09 บาท"
    )
    
    # Q4
    p4 = doc.add_paragraph()
    p4.add_run("คำถามข้อที่ 4: จังหวัดใดมียอดขายสุทธิสูงสุด?\n").font.bold = True
    p4.add_run("• คำตอบ: กรุงเทพมหานคร มียอดขายสุทธิสูงสุด อยู่ที่ ฿2,612,955.88 บาท (154 คำสั่งซื้อ รวม 323 ชิ้น)\n")
    
    # Province Table
    tbl_p = doc.add_table(rows=8, cols=5)
    p_data = [
        ("อันดับ", "จังหวัด (Province)", "จำนวนคำสั่งซื้อ", "จำนวนสินค้า (Qty)", "ยอดขายสุทธิรวม (บาท)"),
        ("1", "กรุงเทพมหานคร", "154", "323", "฿2,612,955.88"),
        ("2", "ขอนแก่น", "110", "225", "฿2,031,943.40"),
        ("3", "ระยอง", "120", "248", "฿1,523,168.61"),
        ("4", "เชียงใหม่", "104", "206", "฿1,477,338.01"),
        ("5", "ภูเก็ต", "86", "164", "฿1,427,388.73"),
        ("6", "ชลบุรี", "86", "171", "฿1,151,249.46"),
        ("รวม", "รวมทั้งสิ้น", "660", "1,337", "฿10,224,044.09")
    ]
    for r_idx, row_vals in enumerate(p_data):
        for c_idx, val in enumerate(row_vals):
            tbl_p.cell(r_idx, c_idx).paragraphs[0].text = val
    format_custom_table(tbl_p, [0.8, 1.8, 1.2, 1.2, 1.5], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Q5
    p5 = doc.add_paragraph()
    p5.add_run("คำถามข้อที่ 5: หมวดสินค้าใดมียอดขายสุทธิสูงสุด?\n").font.bold = True
    p5.add_run("• คำตอบ: หมวดหมู่ Smartphone มียอดขายสุทธิสูงสุด อยู่ที่ ฿3,092,117.34 บาท (178 คำสั่งซื้อ รวม 384 ชิ้น)\n")
    
    # Category Table
    tbl_c = doc.add_table(rows=6, cols=5)
    c_data = [
        ("อันดับ", "หมวดหมู่สินค้า (Category)", "จำนวนคำสั่งซื้อ", "จำนวนสินค้า (Qty)", "ยอดขายสุทธิรวม (บาท)"),
        ("1", "Smartphone", "178", "384", "฿3,092,117.34"),
        ("2", "Accessory", "180", "338", "฿2,710,582.77"),
        ("3", "Notebook", "161", "324", "฿2,221,495.49"),
        ("4", "Smart Home", "141", "291", "฿2,199,848.49"),
        ("รวม", "รวมทั้งสิ้น", "660", "1,337", "฿10,224,044.09")
    ]
    for r_idx, row_vals in enumerate(c_data):
        for c_idx, val in enumerate(row_vals):
            tbl_c.cell(r_idx, c_idx).paragraphs[0].text = val
    format_custom_table(tbl_c, [0.8, 1.8, 1.2, 1.2, 1.5], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Q6
    p6 = doc.add_paragraph()
    p6.add_run("คำถามข้อที่ 6: หากสลับลำดับ merge ก่อน cleaning ผลลัพธ์หรือความเชื่อมั่นของข้อมูลเปลี่ยนอย่างไร?\n").font.bold = True
    add_callout(
        doc,
        [
            "เกิด Cartesian Explosion (Double Counting): ในลูกค้า CRM และ Payments มีข้อมูลซ้ำ การ Merge ก่อน Deduplicate จะทำให้ความสัมพันธ์ M:1 กลายเป็น M:N คำสั่งซื้อจะแตกแถวซ้ำซ้อน ทำให้ยอดขายรวมบวมเกินจริง",
            "ข้อมูลสูญหายจาก Unstandardized Keys (False Unmatched): คีย์ที่มี Whitespace เช่น 'C0001 ' หรือตัวพิมพ์เล็กใหญ่ไม่ตรงกัน หากไม่ Clean ก่อนจะ Merge ไม่ติดและถูกคัดทิ้งอย่างผิดพลาด",
            "การจัดกลุ่ม Dimension ผิดเพี้ยน: หากไม่ Standardize ชื่อจังหวัดก่อน Merge การ Group By จะแยก 'Bangkok', 'กรุงเทพมหานคร', และ 'กทม.' ออกเป็นคนละกลุ่ม ทำให้รายงานผิดพลาด",
            "สูญเสีย Data Lineage & Traceability: การรวมข้อมูลสกปรกเข้าไปก่อนจะทำให้ยากต่อการระบุว่าข้อผิดพลาดเกิดจาก Source ระบบใด"
        ],
        title="บทวิเคราะห์เชิงวิศวกรรมข้อมูล (Data Engineering Insight)"
    )
    
    # -------------------------------------------------------------------------
    # 4. ส่วนท้าทายพิเศษ (Bonus Challenge: Data Validation Framework)
    # -------------------------------------------------------------------------
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. ส่วนท้าทายพิเศษ (Bonus Challenge: Automated Validation Framework)")
    r_h4.font.bold = True
    r_h4.font.size = Pt(13.5)
    r_h4.font.color.rgb = RGBColor(30, 58, 138)
    
    p_ch = doc.add_paragraph(
        "ในไฟล์ starter.py ได้มีการสร้างฟังก์ชัน validate_data(df_fact, df_cust, df_prod) เพื่อตรวจสอบคุณภาพข้อมูลโดยอัตโนมัติ "
        "โดยใช้คำสั่ง assert ตรวจสอบ 3 มิติหลัก ได้แก่:\n"
        "1) Uniqueness: ตรวจสอบความเป็น Unique Key ของ order_id, customer_id, product_id\n"
        "2) Referential Integrity: ตรวจสอบ Foreign Key ทุกรายการใน fact_sales ให้พบใน Dimension Master Data 100%\n"
        "3) Value Ranges: ตรวจสอบ quantity > 0, unit_price > 0, 0 <= discount <= 1 และ net_sales >= 0\n"
        "ผลการประเมิน: ทุกเงื่อนไข Assertion ผ่านการตรวจสอบสมบูรณ์ 100%"
    )
    p_ch.paragraph_format.space_after = Pt(6)
    
    # -------------------------------------------------------------------------
    # 5. รายการไฟล์ผลลัพธ์ (Output Deliverables)
    # -------------------------------------------------------------------------
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. สรุปรายการไฟล์ผลลัพธ์ในโฟลเดอร์ output/ (Output Deliverables)")
    r_h5.font.bold = True
    r_h5.font.size = Pt(13.5)
    r_h5.font.color.rgb = RGBColor(30, 58, 138)
    
    tbl_out = doc.add_table(rows=7, cols=4)
    out_rows = [
        ("ลำดับ", "ชื่อไฟล์ CSV (Output File)", "จำนวนแถว", "รายละเอียดของชุดข้อมูล"),
        ("1", "dim_customer.csv", "160", "มิติลูกค้าหลังตัด Duplicate และ Standardize จังหวัด/Email"),
        ("2", "dim_product.csv", "40", "มิติสินค้าและราคามาตรฐาน พร้อม Active Flag"),
        ("3", "fact_sales.csv", "660", "ตารางข้อเท็จจริงยอดขายที่ผ่านการตรวจสอบ Business Rules และชำระเงินสำเร็จ"),
        ("4", "data_quality_report.csv", "96", "รายงาน Audit Log บันทึกทุกความผิดปกติของข้อมูลและ Action Taken"),
        ("5", "summary_by_province.csv", "6", "ตารางสรุปยอดขายสุทธิ จำนวนออเดอร์ และจำนวนสินค้าแยกตามจังหวัด"),
        ("6", "summary_by_category.csv", "4", "ตารางสรุปยอดขายสุทธิ จำนวนออเดอร์ และจำนวนสินค้าแยกตามหมวดหมู่สินค้า")
    ]
    for r_idx, row_vals in enumerate(out_rows):
        for c_idx, val in enumerate(row_vals):
            tbl_out.cell(r_idx, c_idx).paragraphs[0].text = val
    format_custom_table(tbl_out, [0.6, 2.2, 1.0, 2.7], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    
    # Save files
    out_path1 = Path("Week8/TechTrove_Summary_And_Answers.docx")
    out_path2 = Path("Week8/Lab8_Summary_And_Answers.docx")
    doc.save(out_path1)
    doc.save(out_path2)
    print(f"Summary & Answers docx saved to: {out_path1} and {out_path2}")

if __name__ == "__main__":
    generate_summary_and_answers_document()
