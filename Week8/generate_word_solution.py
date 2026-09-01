import sys
import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1E3A8A"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_table(table, col_widths, col_alignments, header_bg="1E3A8A", alt_bg="F8FAFC"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)
        bg_color = header_bg if is_header else (alt_bg if row_idx % 2 == 1 else "FFFFFF")
        
        # Prevent row split across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if is_header:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[col_idx])
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            for p in cell.paragraphs:
                p.alignment = col_alignments[col_idx]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
                    else:
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(30, 41, 59)

def add_callout_box(doc, text_list, title="ข้อสังเกตสำคัญ (Key Note)", bg_hex="EFF6FF", border_hex="3B82F6"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.font.bold = True
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(30, 58, 138)
    
    for item in text_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        r = p_item.add_run(f"• {item}")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def create_solution_document():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Cordia New'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(30, 41, 59)
    
    # ----------------------------------------------------
    # TITLE SECTION
    # ----------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("รายงานผลการทดลอง: Data Integration Pipeline (Week 8)")
    run_title.font.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("TechTrove E-Commerce: จากข้อมูลดิบหลายระบบสู่ข้อมูลพร้อมวิเคราะห์\nรายวิชา: Data Warehousing Concepts and Design")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # ----------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. บทสรุปผู้บริหาร (Executive Summary)")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(30, 58, 138)
    
    p_exec = doc.add_paragraph(
        "การทดลองนี้เป็นการออกแบบและพัฒนา Data Integration & ETL Pipeline ด้วยภาษา Python และไลบรารี Pandas "
        "เพื่อรวบรวมข้อมูลคำสั่งซื้อจากระบบธุรกรรม (Orders ม.ค. และ ก.พ. 2026), ข้อมูลลูกค้าจากระบบ CRM, "
        "ข้อมูลสินค้าจาก Product Master และเหตุการณ์การชำระเงินจาก Payment Gateway ที่มีโครงสร้างแตกต่างกัน "
        "พร้อมทั้งทำการแก้ไขปัญหา Schema Drift, การทำความสะอาดข้อมูล (Data Cleansing), การตรวจสอบ Referential Integrity, "
        "และการบังคับใช้ Business Rules เพื่อสร้างชุดข้อมูลสำหรับ Data Warehouse ที่ถูกต้อง แม่นยำ และตรวจสอบย้อนกลับได้"
    )
    p_exec.paragraph_format.space_after = Pt(8)
    
    # Summary Metrics Table
    tbl_sum = doc.add_table(rows=6, cols=2)
    sum_data = [
        ("ตัวชี้วัดหลัก (Key Pipeline Metrics)", "ผลลัพธ์ที่ได้จากการประมวลผล (Value)"),
        ("จำนวนคำสั่งซื้อดิบทั้งหมด (Raw Combined Orders)", "752 แถว (ม.ค. 361 แถว + ก.พ. 391 แถว)"),
        ("จำนวนคำสั่งซื้อหลังตัดข้อมูลซ้ำ (Deduplicated Orders)", "750 แถว (พบ Duplicate 2 แถว)"),
        ("จำนวนรายการขายที่สมบูรณ์และถูกต้อง (fact_sales)", "660 ธุรกรรม (อัตราผ่านเกณฑ์ 87.77%)"),
        ("ยอดขายสุทธิรวมทั้งสิ้น (Total Net Sales Revenue)", "฿10,224,044.09 บาท"),
        ("จำนวนเหตุการณ์ Data Quality ที่ถูกบันทึก (DQ Log)", "96 รายการ (บันทึกใน data_quality_report.csv)"),
    ]
    for r_idx, row_vals in enumerate(sum_data):
        for c_idx, val in enumerate(row_vals):
            tbl_sum.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_sum, [3.2, 3.3], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # 2. DATA PROFILING
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. ผลการสำรวจและตรวจสอบคุณภาพข้อมูลดิบ (Data Profiling)")
    r_h2.font.bold = True
    r_h2.font.size = Pt(14)
    r_h2.font.color.rgb = RGBColor(30, 58, 138)
    
    tbl_prof = doc.add_table(rows=6, cols=4)
    prof_data = [
        ("ชุดข้อมูล (Dataset)", "ชนิดไฟล์", "ขนาดดิบ", "ปัญหาคุณภาพข้อมูลที่ตรวจพบ (Data Quality Issues)"),
        ("orders_2026_01.csv", "CSV", "361 × 8", "• Order ID ซ้ำ 1 แถว (ORD000056)\n• Unit price เป็นค่าว่าง 1 แถว (ORD000044)\n• Quantity ติดลบ 1 แถว (ORD000008, qty=-1)"),
        ("orders_2026_02.csv", "CSV", "391 × 8", "• Schema Drift: ชื่อคอลัมน์ ordered_at, qty, discount_pct\n• วันที่รูปแบบ DD/MM/YYYY HH:MM\n• ส่วนลดเป็นสตริงเปอร์เซ็นต์ (เช่น '5%')\n• Order ID ซ้ำ 1 แถว (ORD000416), Unit price ว่าง 1 แถว, Qty ติดลบ 1 แถว"),
        ("customers_crm.csv", "CSV", "163 × 5", "• Customer ID ซ้ำ 3 รายการ (C0012, C0045, C0088)\n• Email เป็นค่าว่าง 5 แถว\n• ชื่อจังหวัดมี 14 รูปแบบที่ไม่เป็นมาตรฐาน (ไทย/อังกฤษ/ตัวย่อ/เว้นวรรค)"),
        ("product_master.xlsx", "Excel", "40 × 5", "• มีข้อมูลสินค้า 40 รายการครบถ้วน ไม่มีค่าว่างหรือรหัสซ้ำ\n• มี Active Flag (Y/N)"),
        ("payments.json", "JSON", "752 events", "• ข้อมูลมีโครงสร้างแบบ Nested JSON\n• Order ID ซ้ำ 1 รายการ (ORD000101)\n• สถานะชำระเงินไม่สำเร็จ: FAILED 47 รายการ, REFUNDED 18 รายการ")
    ]
    for r_idx, row_vals in enumerate(prof_data):
        for c_idx, val in enumerate(row_vals):
            tbl_prof.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_prof, [1.5, 0.8, 0.9, 3.3], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # 3. IMPLEMENTATION STEPS
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. ขั้นตอนการแปลงและทำความสะอาดข้อมูล (ETL Implementation)")
    r_h3.font.bold = True
    r_h3.font.size = Pt(14)
    r_h3.font.color.rgb = RGBColor(30, 58, 138)
    
    p_step1 = doc.add_paragraph()
    p_step1.add_run("3.1 การจัดแนวสคีมาและรวมข้อมูลคำสั่งซื้อ (Schema Alignment & Combine Orders)\n").font.bold = True
    p_step1.add_run(
        "• ปรับเปลี่ยนชื่อคอลัมน์ของเดือน ก.พ. ให้ตรงกับ ม.ค.: ordered_at -> order_date, qty -> quantity, discount_pct -> discount\n"
        "• แปลงรูปแบบส่วนลด: ตัดเครื่องหมาย '%' และหารด้วย 100 เช่น '5%' -> 0.05 (ชนิดข้อมูล float)\n"
        "• แปลงรูปแบบวันเวลา: แปลงสตริงวันที่ทั้งสองเดือนให้อยู่ในรูป Pandas datetime64 มาตรฐาน\n"
        "• ผสานรวมข้อมูล: ใช้คำสั่ง pd.concat([df_jan, df_feb], ignore_index=True) ได้ผลลัพธ์คำสั่งซื้อรวม 752 แถว"
    )
    
    p_step2 = doc.add_paragraph()
    p_step2.add_run("3.2 การทำความสะอาดข้อมูลมิติ (Dimension Cleansing & Standardization)\n").font.bold = True
    p_step2.add_run(
        "• ลูกค้า (Customers CRM): ลบแถวซ้ำโดยเลือกเก็บแถวล่าสุด (keep='last') เหลือ 160 ลูกค้า, ทำความสะอาด Email ด้วย .str.strip().str.lower()\n"
        "• การจัดมาตรฐานชื่อจังหวัด (Province Normalization): ทำการ Mapping ชื่อจังหวัด 14 รูปแบบ สู่ 6 จังหวัดมาตรฐาน ได้แก่:\n"
        "   - 'กรุงเทพมหานคร', 'Bangkok', 'กทม.' -> กรุงเทพมหานคร\n"
        "   - 'ชลบุรี', 'Chonburi', 'ชลบุรี ' -> ชลบุรี\n"
        "   - 'ระยอง', 'Rayong' -> ระยอง\n"
        "   - 'ขอนแก่น', 'ขอนเเก่น' (สระเอสองตัว) -> ขอนแก่น\n"
        "   - 'เชียงใหม่', 'Chiang Mai' -> เชียงใหม่\n"
        "   - 'ภูเก็ต', 'Phuket' -> ภูเก็ต\n"
        "• การชำระเงิน (Payments): แตก Nested JSON และตัด Duplicate order_id (752 -> 751 แถว)"
    )
    
    p_step3 = doc.add_paragraph()
    p_step3.add_run("3.3 การผสานรวมและการตรวจสอบกฎทางธุรกิจ (Integration & Business Validation)\n").font.bold = True
    p_step3.add_run(
        "• ตัดข้อมูลคำสั่งซื้อซ้ำซ้อน: drop_duplicates(subset=['order_id'], keep='last') (752 -> 750 แถว)\n"
        "• เชื่อมโยงข้อมูลด้วย Left Merge: orders -> customers (m:1), orders -> products (m:1), orders -> payments (1:1)\n"
        "• บังคับใช้ Business Rules:\n"
        "   1) quantity > 0 (คัดทิ้ง 2 แถวที่มี qty <= 0)\n"
        "   2) unit_price > 0 และไม่เป็นค่าว่าง (คัดทิ้ง 2 แถวที่มีราคาว่าง)\n"
        "   3) 0 <= discount <= 1\n"
        "   4) Referential Integrity: customer_id และ product_id ต้องปรากฏใน Master Data (คัดทิ้ง 22 ลูกค้า และ 2 สินค้าที่ไม่พบ)\n"
        "   5) สถานะการชำระเงินต้องเป็น 'PAID' เท่านั้น (คัดทิ้ง 46 FAILED และ 18 REFUNDED)\n"
        "• คำนวณ Net Sales: net_sales = quantity * unit_price * (1 - discount)"
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # 4. DATA QUALITY FUNNEL
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. รายงานคุณภาพข้อมูลและ Data Quality Funnel")
    r_h4.font.bold = True
    r_h4.font.size = Pt(14)
    r_h4.font.color.rgb = RGBColor(30, 58, 138)
    
    tbl_funnel = doc.add_table(rows=6, cols=4)
    funnel_data = [
        ("ขั้นตอน (Funnel Stage)", "จำนวนคงเหลือ (Rows)", "จำนวนที่คัดออก (Filtered)", "เหตุผล / คำอธิบาย"),
        ("1. Raw Orders Extracted", "752", "-", "คำสั่งซื้อดิบ ม.ค. (361) + ก.พ. (391)"),
        ("2. Deduplicated Orders", "750", "2", "ลบคำสั่งซื้อซ้ำ (ORD000056, ORD000416)"),
        ("3. Referential Integrity Check", "726", "24", "ไม่พบใน CRM 22 แถว, ไม่พบใน Product Master 2 แถว"),
        ("4. Business Rules Validation", "724", "2", "unit_price เป็น Null 2 แถว, quantity <= 0 อีก 2 แถว (ซ้อนทับกับสถานะชำระเงิน)"),
        ("5. Successful Paid Sales (fact_sales)", "660", "64", "ชำระเงินไม่สำเร็จ (FAILED: 46 รายการ, REFUNDED: 18 รายการ)")
    ]
    for r_idx, row_vals in enumerate(funnel_data):
        for c_idx, val in enumerate(row_vals):
            tbl_funnel.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_funnel, [2.2, 1.1, 1.2, 2.0], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # 5. BUSINESS ANALYSIS QUESTIONS (1-6)
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    r_h5 = h5.add_run("5. คำตอบคำถามวิเคราะห์เชิงธุรกิจ (Business Analysis Questions)")
    r_h5.font.bold = True
    r_h5.font.size = Pt(14)
    r_h5.font.color.rgb = RGBColor(30, 58, 138)
    
    # Question 1
    q1 = doc.add_paragraph()
    q1.add_run("คำถามข้อที่ 1: หลังรวมไฟล์ orders มีจำนวนแถวเท่าใด และเหลือกี่แถวหลังลบ duplicate?\n").font.bold = True
    q1.add_run(
        "• จำนวนแถวหลังรวมไฟล์ Orders (ม.ค. 361 แถว + ก.พ. 391 แถว): รวมทั้งสิ้น 752 แถว\n"
        "• จำนวนแถวที่ถูกลบเนื่องจากซ้ำซ้อน: 2 แถว (ได้แก่ ORD000056 ใน ม.ค. และ ORD000416 ใน ก.พ.)\n"
        "• จำนวนแถวคงเหลือหลังลบ Duplicate: 750 แถว (โดยใช้กฎ subset=['order_id'], keep='last')"
    )
    
    # Question 2
    q2 = doc.add_paragraph()
    q2.add_run("คำถามข้อที่ 2: มีแถวที่ customer_id หรือ product_id ไม่พบใน Master Data อย่างละกี่แถว?\n").font.bold = True
    q2.add_run(
        "• customer_id ที่ไม่พบใน CRM Master Data: มีจำนวน 22 แถว (ได้แก่รหัส C0161, C0162, C0163, C0164, C0165 ซึ่งเป็นลูกค้านอกระบบ)\n"
        "• product_id ที่ไม่พบใน Product Master: มีจำนวน 2 แถว (ได้แก่รหัสสินค้า P999 ในคำสั่งซื้อ ORD000022 และ ORD000382)"
    )
    
    # Question 3
    q3 = doc.add_paragraph()
    q3.add_run("คำถามข้อที่ 3: มียอดขายที่ใช้ได้จริงกี่ธุรกรรม และยอดขายสุทธิรวมเท่าใด?\n").font.bold = True
    q3.add_run(
        "• จำนวนธุรกรรมที่ใช้ได้จริงใน Fact Sales (Valid Paid Transactions): 660 ธุรกรรม (คิดเป็น 87.77% ของคำสั่งซื้อ)\n"
        "• ยอดขายสุทธิรวมทั้งสิ้น (Total Net Sales Revenue): ฿10,224,044.09 บาท"
    )
    
    # Question 4
    q4 = doc.add_paragraph()
    q4.add_run("คำถามข้อที่ 4: จังหวัดใดมียอดขายสุทธิสูงสุด?\n").font.bold = True
    q4.add_run("• คำตอบ: กรุงเทพมหานคร มียอดขายสุทธิสูงสุด อยู่ที่ ฿2,612,955.88 บาท (154 คำสั่งซื้อ รวม 323 ชิ้น)\n")
    
    # Table by Province
    tbl_prov = doc.add_table(rows=8, cols=5)
    prov_rows = [
        ("อันดับ", "จังหวัด (Province)", "จำนวนคำสั่งซื้อ", "จำนวนสินค้า (Qty)", "ยอดขายสุทธิรวม (บาท)"),
        ("1", "กรุงเทพมหานคร", "154", "323", "฿2,612,955.88"),
        ("2", "ขอนแก่น", "110", "225", "฿2,031,943.40"),
        ("3", "ระยอง", "120", "248", "฿1,523,168.61"),
        ("4", "เชียงใหม่", "104", "206", "฿1,477,338.01"),
        ("5", "ภูเก็ต", "86", "164", "฿1,427,388.73"),
        ("6", "ชลบุรี", "86", "171", "฿1,151,249.46"),
        ("รวม", "รวมทั้งสิ้น", "660", "1,337", "฿10,224,044.09")
    ]
    for r_idx, row_vals in enumerate(prov_rows):
        for c_idx, val in enumerate(row_vals):
            tbl_prov.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_prov, [0.8, 1.8, 1.2, 1.2, 1.5], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Question 5
    q5 = doc.add_paragraph()
    q5.add_run("คำถามข้อที่ 5: หมวดสินค้าใดมียอดขายสุทธิสูงสุด?\n").font.bold = True
    q5.add_run("• คำตอบ: หมวดหมู่ Smartphone มียอดขายสุทธิสูงสุด อยู่ที่ ฿3,092,117.34 บาท (178 คำสั่งซื้อ รวม 384 ชิ้น)\n")
    
    # Table by Category
    tbl_cat = doc.add_table(rows=6, cols=5)
    cat_rows = [
        ("อันดับ", "หมวดหมู่สินค้า (Category)", "จำนวนคำสั่งซื้อ", "จำนวนสินค้า (Qty)", "ยอดขายสุทธิรวม (บาท)"),
        ("1", "Smartphone", "178", "384", "฿3,092,117.34"),
        ("2", "Accessory", "180", "338", "฿2,710,582.77"),
        ("3", "Notebook", "161", "324", "฿2,221,495.49"),
        ("4", "Smart Home", "141", "291", "฿2,199,848.49"),
        ("รวม", "รวมทั้งสิ้น", "660", "1,337", "฿10,224,044.09")
    ]
    for r_idx, row_vals in enumerate(cat_rows):
        for c_idx, val in enumerate(row_vals):
            tbl_cat.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_cat, [0.8, 1.8, 1.2, 1.2, 1.5], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT])
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Question 6
    q6 = doc.add_paragraph()
    q6.add_run("คำถามข้อที่ 6: หากสลับลำดับ merge ก่อน cleaning ผลลัพธ์หรือความเชื่อมั่นของข้อมูลเปลี่ยนอย่างไร?\n").font.bold = True
    add_callout_box(
        doc,
        [
            "เกิดปรากฏการณ์ Cartesian Explosion (Double Counting): ในลูกค้า CRM และ Payments มีข้อมูลซ้ำ การ Merge ก่อน Deduplicate จะทำให้ความสัมพันธ์ M:1 กลายเป็น M:N คำสั่งซื้อจะแตกแถวซ้ำซ้อน ทำให้ยอดขายรวมบวมเกินจริง",
            "ข้อมูลสูญหายจาก Unstandardized Keys: รหัสที่มี Whitespace หรือตัวพิมพ์เล็กใหญ่ไม่ตรงกัน หากไม่ทำความสะอาดก่อนจะทำให้ Merge ไม่ติด กลายเป็น False Unmatched และถูกคัดทิ้งอย่างไม่ถูกต้อง",
            "การจัดกลุ่ม Dimension ผิดเพี้ยน: หากไม่ Standardize ชื่อจังหวัดก่อน Merge การ Group By จะแยก 'Bangkok', 'กรุงเทพมหานคร', และ 'กทม.' ออกเป็นคนละกลุ่ม ทำให้รายงานของผู้บริหารผิดพลาด",
            "สูญเสีย Data Lineage & Traceability: การรวมข้อมูลที่สกปรกเข้าไปก่อนจะทำให้ยากต่อการระบุว่าข้อผิดพลาดเกิดจาก Source ระบบใด"
        ],
        title="บทวิเคราะห์เชิงวิศวกรรมข้อมูล (Data Engineering Insight: Merge vs Cleaning Order)"
    )
    
    # ----------------------------------------------------
    # 6. BONUS CHALLENGE & VALIDATION
    # ----------------------------------------------------
    h6 = doc.add_heading(level=1)
    r_h6 = h6.add_run("6. ส่วนท้าทายพิเศษ (Bonus Challenge: Data Validation Framework)")
    r_h6.font.bold = True
    r_h6.font.size = Pt(14)
    r_h6.font.color.rgb = RGBColor(30, 58, 138)
    
    p_chall = doc.add_paragraph(
        "ในการพัฒนา Pipeline ได้มีการสร้างฟังก์ชัน validate_data(df_fact, df_cust, df_prod) เพื่อตรวจสอบคุณภาพข้อมูลโดยอัตโนมัติ "
        "โดยใช้คำสั่ง assert ตรวจสอบ 3 มิติหลัก ได้แก่:\n"
        "1. Uniqueness Assertion: ตรวจสอบความเป็น Unique Key ของ order_id, customer_id, product_id\n"
        "2. Referential Integrity Assertion: ตรวจสอบว่า Foreign Key ทุกตัวใน Fact ปรากฏใน Dimension Tables จริง 100%\n"
        "3. Value Range Assertions: ตรวจสอบ quantity > 0, unit_price > 0, 0 <= discount <= 1 และ net_sales >= 0"
    )
    p_chall.paragraph_format.space_after = Pt(8)
    
    # ----------------------------------------------------
    # 7. OUTPUT DELIVERABLES
    # ----------------------------------------------------
    h7 = doc.add_heading(level=1)
    r_h7 = h7.add_run("7. รายการไฟล์ผลลัพธ์ที่ส่งมอบ (Output Deliverables)")
    r_h7.font.bold = True
    r_h7.font.size = Pt(14)
    r_h7.font.color.rgb = RGBColor(30, 58, 138)
    
    tbl_out = doc.add_table(rows=7, cols=4)
    out_rows = [
        ("ลำดับ", "ชื่อไฟล์ (Output File)", "จำนวนแถว", "รายละเอียดของชุดข้อมูล"),
        ("1", "dim_customer.csv", "160", "มิติลูกค้าหลังตัด Duplicate และ Standardize จังหวัด/Email"),
        ("2", "dim_product.csv", "40", "มิติสินค้าและราคามาตรฐาน พร้อม Active Flag"),
        ("3", "fact_sales.csv", "660", "ตารางข้อเท็จจริงยอดขายที่ผ่านการตรวจสอบ Business Rules และชำระเงินสำเร็จ"),
        ("4", "data_quality_report.csv", "96", "รายงาน Audit Log บันทึกทุกความผิดปกติของข้อมูลและ Action Taken"),
        ("5", "summary_by_province.csv", "6", "ตารางสรุปยอดขายสุทธิ จำนวนออเดอร์ และจำนวนสินค้าแยกตามจังหวัด"),
        ("6", "summary_by_category.csv", "4", "ตารางสรุปยอดขายสุทธิ จำนวนออเดอร์ และจำนวนสินค้าแยกตามหมวดหมู่สินค้า")
    ]
    for r_idx, row_vals in enumerate(out_rows):
        for c_idx, val in enumerate(row_vals):
            tbl_out.cell(r_idx, c_idx).paragraphs[0].text = val
    format_table(tbl_out, [0.6, 2.2, 1.0, 2.7], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    
    # Save document
    output_path = Path(__file__).parent / "Solution.docx"
    doc.save(output_path)
    print(f"Word document successfully saved to: {output_path}")

if __name__ == "__main__":
    create_solution_document()
